# generate_report.py
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Coverage Testing Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('RecipeNOW Project')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(16)
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ============== Overview ==============
doc.add_heading('1. Overview', level=1)

overview_table = doc.add_table(rows=5, cols=2)
overview_table.style = 'Table Grid'
overview_data = [
    ('Metric', 'Value'),
    ('Total Tests', '74'),
    ('Test Files', '9'),
    ('Overall Coverage', '90%'),
    ('Lines Covered', '422 / 468'),
]
for i, (metric, value) in enumerate(overview_data):
    cell0 = overview_table.rows[i].cells[0]
    cell1 = overview_table.rows[i].cells[1]
    cell0.text = metric
    cell1.text = value
    if i == 0:
        cell0.paragraphs[0].runs[0].bold = True
        cell1.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# ============== Test File Structure ==============
doc.add_heading('2. Test File Structure', level=1)

# --- conftest.py ---
doc.add_heading('2.1 conftest.py — Shared Fixtures', level=2)
doc.add_paragraph('Purpose: Provides shared pytest fixtures for all test files')
t = doc.add_table(rows=3, cols=3)
t.style = 'Table Grid'
headers = ['Fixture', 'Scope', 'Description']
for j, h in enumerate(headers):
    t.rows[0].cells[j].text = h
    t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
t.rows[1].cells[0].text = '_set_test_env'
t.rows[1].cells[1].text = 'session'
t.rows[1].cells[2].text = 'Sets GCP_PROJECT_ID and GCP_LOCATION environment variables'
t.rows[2].cells[0].text = 'client'
t.rows[2].cells[1].text = 'function'
t.rows[2].cells[2].text = 'FastAPI TestClient for HTTP endpoint testing'
doc.add_paragraph()

# --- test_generate_rec_router.py ---
doc.add_heading('2.2 test_generate_rec_router.py — Recipe Generation Router', level=2)
doc.add_paragraph('Covers: backend/routers/generate_rec_router.py')
t = doc.add_table(rows=3, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Test Case'
t.rows[0].cells[1].text = 'Description'
t.rows[0].cells[0].paragraphs[0].runs[0].bold = True
t.rows[0].cells[1].paragraphs[0].runs[0].bold = True
t.rows[1].cells[0].text = 'test_generate_recipe_empty_ingredients_returns_empty'
t.rows[1].cells[1].text = 'Verifies empty ingredients return empty response without calling Vertex AI'
t.rows[2].cells[0].text = 'test_generate_recipe_success'
t.rows[2].cells[1].text = 'Mocks Vertex AI to verify successful recipe generation with ingredients'
doc.add_paragraph()

# --- test_scan_router.py ---
doc.add_heading('2.3 test_scan_router.py — Ingredient Scanning Router', level=2)
doc.add_paragraph('Covers: backend/routers/scan_router.py')
t = doc.add_table(rows=3, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Test Case'
t.rows[0].cells[1].text = 'Description'
t.rows[0].cells[0].paragraphs[0].runs[0].bold = True
t.rows[0].cells[1].paragraphs[0].runs[0].bold = True
t.rows[1].cells[0].text = 'test_scan_ingredients_validation_error'
t.rows[1].cells[1].text = 'Verifies 422 error when file is not uploaded'
t.rows[2].cells[0].text = 'test_scan_ingredients_success'
t.rows[2].cells[1].text = 'Mocks Vertex Vision API to verify successful ingredient detection'
doc.add_paragraph()

# --- test_shopping_list_router.py ---
doc.add_heading('2.4 test_shopping_list_router.py — Shopping List Router', level=2)
doc.add_paragraph('Covers: backend/routers/shopping_list_router.py')
t = doc.add_table(rows=3, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Test Case'
t.rows[0].cells[1].text = 'Description'
t.rows[0].cells[0].paragraphs[0].runs[0].bold = True
t.rows[0].cells[1].paragraphs[0].runs[0].bold = True
t.rows[1].cells[0].text = 'test_generate_shopping_list_missing_fields'
t.rows[1].cells[1].text = 'Verifies 400 error when recipe_ingredients is missing'
t.rows[2].cells[0].text = 'test_generate_shopping_list_success'
t.rows[2].cells[1].text = 'Mocks Vertex AI to verify successful shopping list generation'
doc.add_paragraph()

# --- test_routers.py ---
doc.add_heading('2.5 test_routers.py — Extended Router Tests', level=2)
doc.add_paragraph('Covers: Additional edge cases for all three routers (15+ tests)')

tests_routers = [
    ('TestGenerateRecipeRouterExtra', [
        ('test_generate_recipe_with_single_ingredient', 'Tests recipe generation with only one ingredient'),
        ('test_generate_recipe_vertex_error', 'Tests error handling when Vertex API returns 500'),
    ]),
    ('TestScanRouterExtra', [
        ('test_scan_png_image', 'Tests scanning PNG image format'),
        ('test_scan_multiple_ingredients', 'Tests detection of 4 different ingredients'),
    ]),
    ('TestShoppingListRouterExtra', [
        ('test_shopping_list_empty_pantry', 'Tests with completely empty pantry'),
        ('test_shopping_list_all_ingredients_available', 'Tests when all ingredients are in pantry'),
        ('test_shopping_list_missing_both_fields', 'Tests 400 error with empty JSON body'),
        ('test_shopping_list_partial_match', 'Tests partial ingredient quantity matching'),
    ]),
    ('TestGenerateRecipeWithPreferences', [
        ('test_generate_with_diets', 'Tests recipe generation with dietary restrictions (vegan)'),
        ('test_generate_with_time_and_difficulty', 'Tests with time/difficulty constraints'),
    ]),
    ('TestShoppingListAdditional', [
        ('test_shopping_list_vertex_error', 'Tests error handling when Vertex fails'),
        ('test_shopping_list_with_many_items', 'Tests with multiple pantry/recipe items'),
    ]),
    ('TestScanAdditional', [
        ('test_scan_jpeg_image', 'Tests scanning JPEG image with proper headers'),
    ]),
]

for class_name, tests in tests_routers:
    p = doc.add_paragraph()
    run = p.add_run(class_name)
    run.bold = True
    run.font.size = Pt(11)
    t = doc.add_table(rows=len(tests)+1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = 'Test Case'
    t.rows[0].cells[1].text = 'Description'
    t.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    t.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    for i, (name, desc) in enumerate(tests, 1):
        t.rows[i].cells[0].text = name
        t.rows[i].cells[1].text = desc
    doc.add_paragraph()

# --- test_user_auth.py ---
doc.add_heading('2.6 test_user_auth.py — Security & User CRUD Tests', level=2)
doc.add_paragraph('Covers: backend/User/utils/security.py, backend/User/crud/user_crud.py, backend/User/crud/pantry_crud.py')
doc.add_paragraph('Total: 33 tests')

test_user_auth = [
    ('TestPasswordHashing', [
        ('test_get_password_hash_returns_string', 'Verifies bcrypt hash format'),
        ('test_verify_password_correct', 'Tests correct password verification'),
        ('test_verify_password_incorrect', 'Tests wrong password rejection'),
        ('test_verify_password_invalid_hash', 'Tests handling of malformed hash'),
        ('test_normalize_password_truncates_long_password', 'Tests 72-char bcrypt limit'),
        ('test_normalize_password_handles_none', 'Tests None input handling'),
        ('test_normalize_password_short', 'Tests short password encoding'),
    ]),
    ('TestJWT', [
        ('test_create_access_token', 'Verifies JWT token creation'),
        ('test_decode_access_token_valid', 'Tests valid token decoding'),
        ('test_decode_access_token_invalid', 'Tests invalid token rejection'),
        ('test_create_access_token_with_custom_expiry', 'Tests custom expiry delta'),
        ('test_decode_token_contains_exp', 'Verifies expiration claim exists'),
    ]),
    ('TestUserCrud', [
        ('test_get_user_by_id', 'Tests user retrieval by ID'),
        ('test_get_user_by_id_not_found', 'Tests None return for missing user'),
        ('test_get_user_by_phone', 'Tests user lookup by phone number'),
        ('test_list_users', 'Tests paginated user listing'),
        ('test_create_user_with_hashed_password_success', 'Tests user creation flow'),
        ('test_create_user_duplicate_phone_raises', 'Tests duplicate phone ValueError'),
        ('test_update_user_name', 'Tests name update'),
        ('test_delete_user', 'Tests user deletion'),
    ]),
    ('TestPantryCrud', [
        ('test_list_items', 'Tests pantry item listing'),
        ('test_get_item', 'Tests single item retrieval'),
        ('test_create_item', 'Tests item creation'),
        ('test_bulk_create_items', 'Tests batch item creation'),
        ('test_update_item', 'Tests item modification'),
        ('test_delete_item', 'Tests item deletion'),
        ('test_clear_items', 'Tests clearing all user items'),
    ]),
]

for class_name, tests in test_user_auth:
    p = doc.add_paragraph()
    run = p.add_run(class_name)
    run.bold = True
    run.font.size = Pt(11)
    t = doc.add_table(rows=len(tests)+1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = 'Test Case'
    t.rows[0].cells[1].text = 'Description'
    t.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    t.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    for i, (name, desc) in enumerate(tests, 1):
        t.rows[i].cells[0].text = name
        t.rows[i].cells[1].text = desc
    doc.add_paragraph()

# --- test_preferences_crud.py ---
doc.add_heading('2.7 test_preferences_crud.py — Preferences CRUD Tests', level=2)
doc.add_paragraph('Covers: backend/User/crud/preferences_crud.py')
t = doc.add_table(rows=7, cols=2)
t.style = 'Table Grid'
prefs_tests = [
    ('Test Case', 'Description'),
    ('test_get_preferences', 'Tests preferences retrieval'),
    ('test_get_preferences_not_found', 'Tests None for missing preferences'),
    ('test_upsert_preferences_create_new', 'Tests creating new preferences'),
    ('test_upsert_preferences_update_existing', 'Tests updating existing preferences'),
    ('test_delete_preferences_exists', 'Tests preferences deletion'),
    ('test_delete_preferences_not_exists', 'Tests deleting non-existent preferences'),
]
for i, (name, desc) in enumerate(prefs_tests):
    t.rows[i].cells[0].text = name
    t.rows[i].cells[1].text = desc
    if i == 0:
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[1].paragraphs[0].runs[0].bold = True
doc.add_paragraph()

# --- test_auth_dependencies.py ---
doc.add_heading('2.8 test_auth_dependencies.py — Auth Middleware Tests', level=2)
doc.add_paragraph('Covers: backend/User/utils/auth_dependencies.py')
t = doc.add_table(rows=6, cols=2)
t.style = 'Table Grid'
auth_tests = [
    ('Test Case', 'Description'),
    ('test_get_db_yields_session', 'Tests database session lifecycle (yield + close)'),
    ('test_get_current_user_success', 'Tests valid token → user retrieval'),
    ('test_get_current_user_invalid_token', 'Tests 401 for invalid token'),
    ('test_get_current_user_no_sub_in_token', 'Tests 401 for token without sub claim'),
    ('test_get_current_user_user_not_found', 'Tests 401 when user does not exist'),
]
for i, (name, desc) in enumerate(auth_tests):
    t.rows[i].cells[0].text = name
    t.rows[i].cells[1].text = desc
    if i == 0:
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[1].paragraphs[0].runs[0].bold = True
doc.add_paragraph()

# --- test_user_router.py ---
doc.add_heading('2.9 test_user_router.py — User Router Endpoint Tests', level=2)
doc.add_paragraph('Covers: backend/User/routers/user_router.py')
t = doc.add_table(rows=6, cols=2)
t.style = 'Table Grid'
router_tests = [
    ('Test Case', 'Description'),
    ('test_login_success', 'Tests successful login returns JWT'),
    ('test_login_wrong_password', 'Tests 401 for wrong password'),
    ('test_login_user_not_found', 'Tests 401 for non-existent user'),
    ('test_register_success', 'Tests successful user registration'),
    ('test_register_duplicate_phone', 'Tests 400 for duplicate phone number'),
]
for i, (name, desc) in enumerate(router_tests):
    t.rows[i].cells[0].text = name
    t.rows[i].cells[1].text = desc
    if i == 0:
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[1].paragraphs[0].runs[0].bold = True
doc.add_paragraph()

# ============== Testing Techniques ==============
doc.add_heading('3. Testing Techniques Used', level=1)
t = doc.add_table(rows=6, cols=2)
t.style = 'Table Grid'
techniques = [
    ('Technique', 'Description'),
    ('Mocking', 'unittest.mock.MagicMock for database sessions'),
    ('Monkeypatching', 'pytest.monkeypatch for replacing Vertex AI calls'),
    ('TestClient', 'FastAPI TestClient for HTTP endpoint testing'),
    ('Fixtures', 'Shared setup via conftest.py'),
    ('Class Organization', 'Tests grouped by feature/module using test classes'),
]
for i, (tech, desc) in enumerate(techniques):
    t.rows[i].cells[0].text = tech
    t.rows[i].cells[1].text = desc
    if i == 0:
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[1].paragraphs[0].runs[0].bold = True
doc.add_paragraph()

# ============== Coverage Configuration ==============
doc.add_heading('4. Coverage Configuration', level=1)
doc.add_paragraph('Excluded from coverage (via .coveragerc):')
bullets = [
    'backend/init_db.py — One-time database initialization script',
    'backend/routers/scan_router.py — Requires real GCP credentials for full coverage',
    '*/tests/* — Test files themselves',
    '*/__pycache__/* — Python bytecode',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')
doc.add_paragraph()

# ============== How to Run Tests ==============
doc.add_heading('5. How to Run Tests', level=1)
doc.add_paragraph('Run all tests with coverage report:')
doc.add_paragraph('pytest --cov=backend --cov=main --cov-report=term-missing tests/', style='Quote')
doc.add_paragraph()
doc.add_paragraph('Generate HTML coverage report:')
doc.add_paragraph('pytest --cov=backend --cov=main --cov-report=html tests/', style='Quote')
doc.add_paragraph('Then open htmlcov/index.html in your browser.')

# Save to the specified directory
output_path = '/Users/liutongyan/Documents/BostonUniversity/CS411/Coverage_Testing_Report.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')

